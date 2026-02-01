import os
import sys

# Add backend to path
sys.path.append(os.path.join(os.getcwd(), "app"))

try:
    from docx import Document
    from app.utils.resume_parser import parse_resume_content
except ImportError as e:
    print(f"ImportError: {e}")
    print("Make sure you are running this from i:\\InterviewProject\\backend")
    sys.exit(1)


def create_test_docx(filename):
    doc = Document()
    doc.add_heading("John Doe", 0)
    doc.add_paragraph(
        "Software Engineer with 5 years of experience in Python and React."
    )
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Senior Developer at Tech Corp.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, React, MongoDB")
    doc.save(filename)
    print(f"Created {filename}")


def test_parser():
    filename = "test_resume.docx"
    create_test_docx(filename)

    with open(filename, "rb") as f:
        content = f.read()

    text = parse_resume_content(content, filename)
    print("\n--- Extracted Text ---")
    print(text)
    print("----------------------")

    if "John Doe" in text and "Python" in text:
        print("✅ Parsing SUCCESS")
    else:
        print("❌ Parsing FAILED")

    # Cleanup
    try:
        os.remove(filename)
    except:
        pass


if __name__ == "__main__":
    test_parser()
